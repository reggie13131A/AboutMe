import getpass
import os
import os
from langchain_neo4j import Neo4jGraph
from langchain_experimental.graph_transformers import LLMGraphTransformer
from langchain_community.chat_models import ChatTongyi
from langchain_core.documents import Document
import asyncio
from langchain_neo4j import GraphCypherQAChain, Neo4jGraph
from docx import Document as DocxDocument
import pandas as pd
from langchain_community.llms import HuggingFacePipeline
from langchain_community.chat_models import ChatHuggingFace
import torch
from transformers import BitsAndBytesConfig

# -i https://mirrors.aliyun.com/pypi/simple

os.environ["DASHSCOPE_API_KEY"] = "sk-7ac5a239fc414c79b36968ed1b6b9b0b"

os.environ["NEO4J_URI"] = "bolt://localhost:7687"
os.environ["NEO4J_USERNAME"] = "neo4j"
os.environ["NEO4J_PASSWORD"] = "Pensoon123"

graph = Neo4jGraph()

llm_for_extraction = ChatTongyi(
    model="qwen-max",
    temperature=0,
    max_tokens=2000,
    disable_streaming=True,
)

llm_transformer = LLMGraphTransformer(llm=llm_for_extraction)

def read_xlsx_file(file_path):
    """
    读取xlsx文件并返回文本内容
    """
    try:
        # 读取Excel文件的所有工作表
        excel_file = pd.ExcelFile(file_path)
        all_text = []
        
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            # 添加工作表名称
            all_text.append(f"工作表: {sheet_name}\n")
            # 将DataFrame转换为文本
            all_text.append(df.to_string(index=False))
            all_text.append("\n" + "="*50 + "\n")
        
        return "\n".join(all_text)
    except Exception as e:
        print(f"读取Excel文件时出错: {str(e)}")
        return ""

def read_docx_file(file_path):
    """
    读取docx文件并返回文本内容
    """
    try:
        doc = DocxDocument(file_path)
        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        return '\n'.join(full_text)
    except Exception as e:
        print(f"读取Word文件时出错: {str(e)}")
        return ""

async def process_file(file_path):
    """
    根据文件扩展名处理不同类型的文件
    """
    if file_path.endswith('.docx'):
        text_content = read_docx_file(file_path)
    elif file_path.endswith('.xlsx'):
        text_content = read_xlsx_file(file_path)
    else:
        print(f"不支持的文件类型: {file_path}")
        return []
    
    if not text_content.strip():
        print(f"文件 {file_path} 没有内容")
        return []
    
    # 将文本内容分割为多个文档片段
    documents = []
    lines = text_content.split('\n')
    current_chunk = []
    chunk_size = 500  # 每个文档片段的大致行数
    
    for i, line in enumerate(lines):
        current_chunk.append(line)
        if len(current_chunk) >= chunk_size or i == len(lines) - 1:
            chunk_text = '\n'.join(current_chunk)
            if chunk_text.strip():
                documents.append(Document(
                    page_content=chunk_text,
                    metadata={"source": os.path.basename(file_path), "chunk": len(documents)}
                ))
            current_chunk = []
    
    return documents

async def main():
    try:
        # 处理多个文件
        file_paths = [
            # "/home/admin-ps/Documents/hatAndMask/ChatRobot/datasets/2025.docx",
            "/home/admin-ps/Documents/hatAndMask/ChatRobot/datasets/知识库.xlsx"  # 添加您的Excel文件路径
        ]
        
        all_documents = []
        for file_path in file_paths:
            if os.path.exists(file_path):
                print(f"正在处理文件: {file_path}")
                documents = await process_file(file_path)
                all_documents.extend(documents)
                print(f"从 {file_path} 提取了 {len(documents)} 个文档片段")
            else:
                print(f"文件不存在: {file_path}")
        
        print(f"总共处理了 {len(all_documents)} 个文档片段")
        
        if not all_documents:
            print("没有有效的文档内容可处理")
            return
            
        # 转换为图文档
        print("开始转换文档为图结构...")
        graph_documents = await llm_transformer.aconvert_to_graph_documents(all_documents)
        
        if graph_documents:
            print(f"成功转换 {len(graph_documents)} 个图文档")
            print(f"第一个文档的节点数: {len(graph_documents[0].nodes)}")
            print(f"第一个文档的关系数: {len(graph_documents[0].relationships)}")
            
            # 添加到图数据库
            graph.add_graph_documents(graph_documents)
            print("成功添加到图数据库")
        else:
            print("没有生成图文档")
            
    except Exception as e:
        print(f"处理过程中发生错误: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    asyncio.run(main())
    graph.refresh_schema()
    quantization_config = BitsAndBytesConfig(
        load_in_4bit=True,
        bnb_4bit_compute_dtype=torch.float16
    )
    # print(graph.schema)
    
    # chain = GraphCypherQAChain.from_llm(
    #     ChatTongyi(
    #         model="qwen-max",
    #         temperature=0,
    #         max_tokens=4000
    #     ),
    #     graph=graph,
    #     return_intermediate_steps=True,
    #     allow_dangerous_requests=True
    # )
    llm = HuggingFacePipeline.from_model_id(
        model_id="/home/admin-ps/Documents/hatAndMask/deepseek7b",
        task="text-generation",
        pipeline_kwargs={
            "max_new_tokens": 4000,
            "do_sample": False
        },
        model_kwargs={
            "quantization_config": quantization_config,
            "trust_remote_code": True
        }
    )   

    chain = GraphCypherQAChain.from_llm(
        llm,
        graph=graph,
        validate_cypher= True,
        return_intermediate_steps=True,
        allow_dangerous_requests=True
    )

    result = chain.invoke({"query": "天然气业务有哪些呢"})
    print(f"Intermediate steps: {result['intermediate_steps']}")
    print(f"Final answer: {result['result']}")


